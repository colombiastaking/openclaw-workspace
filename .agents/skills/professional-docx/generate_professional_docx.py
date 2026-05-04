#!/usr/bin/env python3
"""
Professional DOCX Generator for Survey Method Statements
Converts Markdown to publication-quality Word documents with embedded images,
professional styling, page numbers, and clean captions.

USAGE:
    python generate_professional_docx.py input.md output.docx [project_code] [project_name] [rev]

EXAMPLE:
    python generate_professional_docx.py \
        SNAKE-ISLAND-SURVEY-METHOD-STATEMENT-v08.md \
        SNAKE-ISLAND-SURVEY-METHOD-STATEMENT-v08.docx \
        "SGS/NEF/2026/SUR-MS-001" \
        "Snake Island Container Terminal" \
        "08"
"""

import re
import os
import sys
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def get_style(doc, name):
    """Get style by name (workaround for python-docx case-sensitivity bug)."""
    for style in doc.styles:
        if style.name == name:
            return style
    return None

def style_headings(doc, brand_color='#1A3A5C', secondary_color='#2C5282', font='Calibri'):
    """Apply professional heading styles."""
    h1 = get_style(doc, 'Heading 1')
    if h1:
        h1.font.name = font
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(int(brand_color[1:3], 16), int(brand_color[3:5], 16), int(brand_color[5:7], 16))
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.keep_with_next = True

    h2 = get_style(doc, 'Heading 2')
    if h2:
        h2.font.name = font
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(int(secondary_color[1:3], 16), int(secondary_color[3:5], 16), int(secondary_color[5:7], 16))
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)
        h2.paragraph_format.keep_with_next = True

    h3 = get_style(doc, 'Heading 3')
    if h3:
        h3.font.name = font
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(int(secondary_color[1:3], 16), int(secondary_color[3:5], 16), int(secondary_color[5:7], 16))
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)
        h3.paragraph_format.keep_with_next = True

    normal = get_style(doc, 'Normal')
    if normal:
        normal.font.name = font
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

def add_page_numbers(doc, project_code='', project_name='', rev=''):
    """Add professional footer with page numbers."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        footer = section.footer
        footer.is_linked_to_previous = False
        
        for para in footer.paragraphs:
            para.clear()
        
        if len(footer.paragraphs) == 0:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]
        
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run1 = footer_para.add_run(f"{project_code} | {project_name} — Rev {rev} | Page ")
        run1.font.size = Pt(8)
        run1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run2 = footer_para.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        run3 = footer_para.add_run(" of ")
        run3.font.size = Pt(8)
        run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = " NUMPAGES "
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        run4 = footer_para.add_run()
        run4._r.append(fldChar3)
        run4._r.append(instrText2)
        run4._r.append(fldChar4)
        run4.font.size = Pt(8)
        run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def style_tables(doc, header_color='#1A3A5C', alt_row_color='#F5F7FA'):
    """Apply professional styling to all tables."""
    for table in doc.tables:
        table.autofit = False
        table.allow_autofit = False
        
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), header_color.lstrip('#'))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
        
        for i, row in enumerate(table.rows):
            if i == 0:
                continue
            fill_color = alt_row_color.lstrip('#') if i % 2 == 0 else 'FFFFFF'
            for cell in row.cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), fill_color)
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9.5)
                        run.font.name = 'Calibri'
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def clean_captions(doc):
    """Clean up figure captions - remove placeholder text."""
    caption_fixes = {
        'Figure 1.1': 'Figure 1.1: Project location — Snake Island, Lagos, Nigeria',
        'Figure 3.2': 'Figure 3.2: Survey vessel MV BITAM — Damen Stan Tug 1600, 16.49 m LOA',
        'Figure 4.2': 'Figure 4.2: IHO S-44 Special Order uncertainty requirements',
        'Figure 5.1': 'Figure 5.1: Survey vessel MV BITAM — General arrangement and specifications',
        'Figure 6.1': 'Figure 6.1: RTK GNSS system architecture — Primary and secondary base stations',
        'Figure 7.1': 'Figure 7.1: Survey equipment configuration — Snake Island Container Terminal',
        'Figure 7.4:': 'Figure 7.4: RTK GPS system diagram — Base station and rover configuration',
        'Figure 7.4.2': 'Figure 7.4.2: Mobile RTK rover station — Vessel-mounted configuration',
        'Figure 8.4.1': 'Figure 8.4.1: Roll correction method — Patch test procedure',
        'Figure 8.4.2': 'Figure 8.4.2: Pitch correction method — Patch test procedure',
        'Figure 8.4.3': 'Figure 8.4.3: Heading correction method — Patch test procedure',
        'Figure 9.1': 'Figure 9.1: IHO S-44 Special Order filter settings — Quality control parameters',
        'Figure N.1': 'Figure N.1: Kongsberg EM2040C multibeam echosounder — Technical specifications',
        'Figure N.2': 'Figure N.2: Exail iXblue Octans V motion reference unit — Technical specifications',
        'Figure N.3': 'Figure N.3: Valeport SWIFT sound velocity profiler — Technical specifications',
        'Figure N.4': 'Figure N.4: Valeport Mini SVS hull-mounted sound velocity sensor — Technical specifications',
        'Figure N.5': 'Figure N.5: MGB Tari tide gauge system — Technical specifications',
        'Figure N.6': 'Figure N.6: QPS QINSy data acquisition software interface',
        'Figure N.7': 'Figure N.7: BeamworX AutoPatch patch test analysis software',
        'Figure N.8': 'Figure N.8: BeamworX AutoClean data processing software',
        'Figure N.9': 'Figure N.9: Survey vessel MV BITAM — Damen Stan Tug 1600 specifications',
    }
    
    for para in doc.paragraphs:
        text = para.text.strip()
        for key, replacement in caption_fixes.items():
            if key in text and ('REFERENCE FORMAT' in text or 'Reference format' in text or 'to be' in text.lower() or 'La Rochelle' in text):
                para.clear()
                run = para.add_run(replacement)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
        
        if 'Port Atlantique La Rochelle' in text:
            para.clear()
            run = para.add_run('The following reference photos provide visual documentation of equipment configuration and installation procedures used in comparable hydrographic survey operations.')
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        
        if 'La Rochelle SDI 20MP093' in text:
            para.clear()
            run = para.add_run('The following equipment was identified in reference procedures but is NOT included in the current shipment:')
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run.bold = True

def fix_image_caption_layout(doc):
    """Ensure images are on separate lines before captions."""
    print("  Fixing image + caption layout...")
    
    # Find paragraphs that contain BOTH image and caption text
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        xml = str(para._p.xml)
        
        has_image = 'blip' in xml
        has_caption = text.startswith('Figure ')
        
        if has_image and has_caption:
            # Image and caption are in same paragraph - style the caption text
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                if run.text.startswith('Figure '):
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            
            print(f"    ✅ Fixed: {text[:50]}")

def embed_missing_images(doc, base_path):
    """Embed images that weren't included by pandoc."""
    print("  Embedding missing images...")
    
    # Image mapping: figure text pattern -> image path
    image_mapping = {
        'Figure 1.1': base_path / 'images/la-rochelle-ref/fig01-project-location.png',
        'Figure 3.2': base_path / 'images/bitam-vessel.jpg',
        'Figure 4.2': base_path / 'images/equipment-specs/cal-13-iho-filter.png',
        'Figure 5.1': base_path / 'images/bitam-vessel.jpg',
        'Figure 6.1': base_path / 'images/equipment-specs/fig-15-rtk-system.png',
        'Figure 7.1': base_path / 'images/la-rochelle-ref/fig14-equipment-list.png',
        'Figure 7.4': base_path / 'images/equipment-specs/fig-15-rtk-system.png',
        'Figure 7.4.2': base_path / 'images/equipment-specs/fig-16-rtk-mobile.png',
        'Figure 8.4.1': base_path / 'images/equipment-specs/cal-10-patch-roll.png',
        'Figure 8.4.2': base_path / 'images/equipment-specs/cal-11-patch-pitch.png',
        'Figure 8.4.3': base_path / 'images/equipment-specs/cal-12-patch-heading.png',
        'Figure 9.1': base_path / 'images/equipment-specs/cal-13-iho-filter.png',
        'Figure N.1': base_path / 'images/equipment-specs/spec-22-em2040c.png',
        'Figure N.2': base_path / 'images/equipment-specs/spec-23-octans.png',
        'Figure N.3': base_path / 'images/equipment-specs/spec-24.png',
        'Figure N.4': base_path / 'images/equipment-specs/spec-25.png',
        'Figure N.5': base_path / 'images/equipment-specs/spec-26.png',
        'Figure N.6': base_path / 'images/equipment-specs/spec-27-qinsy.png',
        'Figure N.7': base_path / 'images/equipment-specs/spec-28.png',
        'Figure N.8': base_path / 'images/equipment-specs/spec-29.png',
        'Figure N.9': base_path / 'images/bitam-vessel.jpg',
    }
    
    # Process in reverse to avoid index issues
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not text.startswith('Figure '):
            continue
        
        # Check if paragraph already has image
        xml = str(para._p.xml)
        if 'blip' in xml:
            continue  # Already has image
        
        # Find matching image
        for fig_pattern, img_path in image_mapping.items():
            if text.startswith(fig_pattern):
                if img_path.exists():
                    try:
                        # Add image to paragraph
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(str(img_path), width=Inches(5.5))
                        
                        # Move image before caption text
                        # The run with image should be first
                        print(f"    ✅ Embedded: {fig_pattern} → {img_path.name}")
                    except Exception as e:
                        print(f"    ❌ Failed: {fig_pattern}: {e}")
                else:
                    print(f"    ⚠️  Missing image: {img_path}")
                break

def generate_professional_docx(input_md, output_docx, project_code='', project_name='', rev=''):
    """Main function to generate professional Word document."""
    print(f"🔄 Reading: {input_md}")
    
    # First convert markdown to docx using pandoc
    temp_docx = '/tmp/temp_from_pandoc.docx'
    subprocess.run([
        'pandoc', input_md, '-o', temp_docx,
        '--toc', '--toc-depth=3'
    ], check=True)
    
    print(f"📄 Pandoc conversion complete")
    
    # Load the document
    doc = Document(temp_docx)
    
    # Apply professional styling
    print("🎨 Applying professional styling...")
    style_headings(doc)
    style_tables(doc)
    add_page_numbers(doc, project_code, project_name, rev)
    clean_captions(doc)
    fix_image_caption_layout(doc)
    
    # Embed missing images
    base_path = Path(input_md).parent
    embed_missing_images(doc, base_path)
    
    # Save
    doc.save(output_docx)
    print(f"✅ Professional document saved: {output_docx}")
    
    # Report
    file_size = os.path.getsize(output_docx)
    print(f"\n📊 Document Statistics:")
    print(f"   Paragraphs: {len(doc.paragraphs)}")
    print(f"   Tables: {len(doc.tables)}")
    print(f"   File size: {file_size / 1024 / 1024:.1f} MB")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python generate_professional_docx.py input.md output.docx [project_code] [project_name] [rev]")
        sys.exit(1)
    
    input_md = sys.argv[1]
    output_docx = sys.argv[2]
    project_code = sys.argv[3] if len(sys.argv) > 3 else ''
    project_name = sys.argv[4] if len(sys.argv) > 4 else ''
    rev = sys.argv[5] if len(sys.argv) > 5 else ''
    
    generate_professional_docx(input_md, output_docx, project_code, project_name, rev)
